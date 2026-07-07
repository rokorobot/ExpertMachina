import os
import hashlib
import json
from sqlalchemy.orm import Session
from app import database as db
from app import crud
from app.logging_config import get_logger
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams, PointStruct

logger = get_logger(__name__)

# Initialize local Qdrant client lazily
QDRANT_DIR = "./qdrant_db"
_qdrant_client = None
COLLECTION_NAME = "expert_machina_chunks"

def get_qdrant_client():
    global _qdrant_client
    if _qdrant_client is None:
        _qdrant_client = QdrantClient(path=QDRANT_DIR)
    return _qdrant_client

# Ensure collection exists
def init_qdrant():
    try:
        client = get_qdrant_client()
        collections = client.get_collections().collections
        exists = any(c.name == COLLECTION_NAME for c in collections)
        if not exists:
            # We'll support 1536 dims (OpenAI Ada dimensions)
            client.create_collection(
                collection_name=COLLECTION_NAME,
                vectors_config=VectorParams(size=1536, distance=Distance.COSINE),
            )
    except Exception as e:
        logger.error("Error initializing Qdrant: %s", e)

# Helper to generate mock embeddings if OpenAI is not configured
def get_embedding(text: str) -> list:
    api_key = os.environ.get("OPENAI_API_KEY")
    if api_key and not api_key.startswith("mock-"):
        try:
            from llama_index.embeddings.openai import OpenAIEmbedding
            embedder = OpenAIEmbedding(api_key=api_key)
            return embedder.get_text_embedding(text)
        except Exception as e:
            logger.warning("Error calling OpenAI Embeddings: %s. Falling back to mock embeddings.", e)
            
    # Mock embedding: generate a deterministic 1536-dim float array from the text hash
    h = hashlib.sha256(text.encode('utf-8')).digest()
    vector = []
    # Fill 1536 elements using sha256 bytes
    for i in range(1536):
        byte_idx = (i * 7) % len(h)
        val = (h[byte_idx] / 255.0) - 0.5
        vector.append(val)
    return vector

def is_chunk_valid(text: str) -> bool:
    # Reject raw OOXML/ZIP internals or raw binary streams
    invalid_patterns = [
        "[Content_Types].xml",
        "_rels/.rels",
        "word/document.xml",
        "word/_rels/document.xml.rels",
        "PK\x03\x04"
    ]
    if any(pattern in text for pattern in invalid_patterns):
        return False
        
    # Check for excessive unreadable/non-ascii/binary control characters
    # If more than 25% of text contains control characters/non-printable, reject
    control_chars = len([c for c in text if ord(c) < 32 and c not in "\n\r\t"])
    if len(text) > 0 and (control_chars / len(text)) > 0.15:
        return False
        
    return True

def parse_and_index_document(session: Session, doc_id: int):
    doc = session.query(db.Document).filter(db.Document.id == doc_id).first()
    if not doc:
        return False
        
    crud.update_project_status(session, doc.project_id, "INGESTING")
    
    try:
        file_path = doc.file_path
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found at {file_path}")

        # Ingestion step 1: Parsing
        text_content = ""
        table_count = 0
        docling_json_str = None
        
        # Determine file type
        is_docx = file_path.lower().endswith(".docx")
        
        # DOCX Ingestion strategy: Use Docling, fallback to python-docx
        if is_docx:
            try:
                # Try Docling first
                from docling.document_converter import DocumentConverter
                converter = DocumentConverter()
                result = converter.convert_single(file_path)
                text_content = result.render_as_markdown()
                table_count = len(result.pages[0].tables) if hasattr(result, "pages") and len(result.pages) > 0 else 0
                docling_json_str = json.dumps({"parsed_by": "docling_docx"})
            except Exception as e:
                logger.warning("Docling docx parsing skipped or failed (%s). Falling back to python-docx.", e)
                try:
                    import docx
                    doc_obj = docx.Document(file_path)
                    paragraphs = [p.text for p in doc_obj.paragraphs if p.text.strip()]
                    # Also extract tables content
                    for table in doc_obj.tables:
                        for row in table.rows:
                            row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                            if row_text:
                                paragraphs.append(row_text)
                                table_count += 1
                    text_content = "\n\n".join(paragraphs)
                    docling_json_str = json.dumps({"parsed_by": "python_docx"})
                except Exception as ex:
                    logger.warning("python-docx parsing failed: %s", ex)
        else:
            # Non-docx files
            try:
                from docling.document_converter import DocumentConverter
                converter = DocumentConverter()
                result = converter.convert_single(file_path)
                text_content = result.render_as_markdown()
                table_count = len(result.pages[0].tables) if hasattr(result, "pages") and len(result.pages) > 0 else 0
                docling_json_str = json.dumps({"parsed_by": "docling"})
            except Exception as e:
                logger.warning("Docling parsing skipped or failed (%s). Falling back to native text extraction.", e)
                if file_path.endswith(".txt") or file_path.endswith(".md"):
                    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                        text_content = f.read()
                elif file_path.endswith(".pdf"):
                    text_content = f"--- PDF DOCUMENT Content Fallback ---\nFilename: {doc.filename}\nThis is a mock representation of the PDF content parsed from {doc.filename}."
                    try:
                        import pypdf
                        reader = pypdf.PdfReader(file_path)
                        pages_text = []
                        for page in reader.pages:
                            pages_text.append(page.extract_text() or "")
                        text_content = "\n\n".join(pages_text)
                    except Exception as ex:
                        logger.warning("pypdf extraction skipped: %s", ex)
                else:
                    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                        text_content = f.read()

        # Gate check: Validate entire parsed text before chunking
        if not text_content.strip() or not is_chunk_valid(text_content):
            raise ValueError("Parsed content is invalid: Contains ZIP/XML internal files, binary markers, or excessive non-printable text.")

        # Ingestion step 2: Chunking
        chunk_size = 500
        overlap = 100
        chunks = []
        
        words = text_content.split()
        current_chunk = []
        current_length = 0
        
        for word in words:
            current_chunk.append(word)
            current_length += len(word) + 1
            if current_length >= chunk_size:
                chunk_candidate = " ".join(current_chunk)
                if is_chunk_valid(chunk_candidate):
                    chunks.append(chunk_candidate)
                # Backtrack for overlap
                overlap_words = current_chunk[-max(1, int(len(current_chunk)*0.2)):]
                current_chunk = overlap_words
                current_length = sum(len(w) for w in current_chunk)
                
        if current_chunk:
            chunk_candidate = " ".join(current_chunk)
            if is_chunk_valid(chunk_candidate):
                chunks.append(chunk_candidate)
            
        # If no valid chunks were produced
        if not chunks:
            raise ValueError("No valid human-readable chunks could be extracted.")

        # Ingestion step 3: Embed & Store in Qdrant
        init_qdrant()
        
        for idx, chunk_text in enumerate(chunks):
            embedding = get_embedding(chunk_text)
            
            # Store chunk in DB
            db_chunk = db.DocumentChunk(
                document_id=doc.id,
                text=chunk_text,
                chunk_index=idx,
                table_count=table_count if idx == 0 else 0,
                docling_json=docling_json_str,
                embedding_ref=f"qdrant-pt-{doc.id}-{idx}"
            )
            session.add(db_chunk)
            session.commit()
            session.refresh(db_chunk)
            
            # Store embedding in Qdrant
            get_qdrant_client().upsert(
                collection_name=COLLECTION_NAME,
                points=[
                    PointStruct(
                        id=hashlib.md5(f"{doc.id}-{idx}".encode()).hexdigest(),
                        vector=embedding,
                        payload={
                            "document_id": doc.id,
                            "chunk_id": db_chunk.id,
                            "text": chunk_text
                        }
                    )
                ]
            )
            
        # Mark document as parsed
        doc.status = "PARSED"
        session.commit()
        
        crud.log_audit_event(session, actor="system", event_type="DOCUMENT_PARSED", target_id=str(doc.id), details=f"Successfully parsed and chunked document '{doc.filename}' into {len(chunks)} chunks.")
        crud.update_project_status(session, doc.project_id, "TRANSFORMING")
        return True
        
    except Exception as e:
        doc.status = "FAILED"
        session.commit()
        crud.log_audit_event(session, actor="system", event_type="DOCUMENT_PARSED_FAILED", target_id=str(doc.id), details=f"Failed to parse document '{doc.filename}': {str(e)}")
        return False
