import os
import json
import hashlib
import datetime

from sqlalchemy.orm import Session

from app import database as db
from app import crud
from app import ingestion
from app import extraction
from app import policy
from app import revisions
from app.connectors.models import ConnectorItem
from app.connectors.providers.local_folder import LocalFolderProvider

# Enterprise Source Connector - LOCAL_FOLDER (MVP 0.10.0).
#
# Read-only discovery over a local or mounted directory. "Scan now" only:
# recursive walk, extension filter, sha256 dedup, per-file status. No
# watching, no cloud connectors, no credentials. Newly ingested CANDIDATE
# assets may be auto-approved by explicit per-project ApprovalPolicy rules
# (MVP 0.10.2, see policy.py); changed files always wait for a human.
#
# The architectural rule: connector output becomes ORDINARY ExpertMachina
# objects. SourceDocument -> Document -> CANDIDATE KnowledgeAssets -> the
# existing governance pipeline (review queue, conflicts, revisions, inbox).
# No connector-specific approval flow, no special "folder assets".
#
# Follows the v0.9.2a discipline: the scan endpoint returns immediately;
# the heavy work runs as a background task with its own session, updating
# job counters per file (live progress) and writing its own audit events.

SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}
UPLOAD_DIR = "./uploads"


def _provider_for(connector: db.SourceConnector) -> LocalFolderProvider:
    """Provider construction. Only LOCAL_FOLDER exists; a registry is earned
    when a second provider type does (D8). SUPPORTED_EXTENSIONS is the
    framework's parser capability, handed to the provider so enumeration
    filters on the capability/config intersection (seam 1)."""
    return LocalFolderProvider(connector, SUPPORTED_EXTENSIONS)


def discover_files(connector: db.SourceConnector) -> list:
    """Pre-0.11 compatibility surface: the walk now lives in
    LocalFolderProvider.discover(); returns the discovered URIs
    (absolute paths, deterministic order) as plain strings."""
    return [item.uri for item in _provider_for(connector).discover()]


def run_ingestion_job(job_id: int):
    """Background-task entry point: owns its session (the request that
    scheduled the scan has already returned)."""
    session = db.SessionLocal()
    try:
        execute_ingestion_job(session, job_id)
    finally:
        session.close()


def execute_ingestion_job(session: Session, job_id: int, auto_extract: bool = True) -> db.IngestionJob:
    job = session.query(db.IngestionJob).filter(db.IngestionJob.id == job_id).first()
    if not job:
        print(f"Ingestion job {job_id} not found")
        return None
    connector = session.query(db.SourceConnector).filter(
        db.SourceConnector.id == job.connector_id).first()

    provider = _provider_for(connector)

    job.status = "RUNNING"
    job.started_at = datetime.datetime.utcnow()
    session.commit()
    # describe(): provider-specific source context, logged without
    # interpretation. Key order keeps the payload identical to pre-0.11.
    crud.log_audit_event(
        session, actor=f"connector:{connector.name}", event_type="INGESTION_JOB_STARTED",
        target_id=str(job.id),
        details=json.dumps({"connector_id": connector.id, **provider.describe()}))

    try:
        # Reachability is the provider's question; whether ingestion runs
        # remains the framework's.
        provider.validate()

        items = provider.discover()
        job.files_discovered = len(items)
        session.commit()

        os.makedirs(UPLOAD_DIR, exist_ok=True)
        seen_hashes_this_job = set()

        for item in items:
            outcome = _ingest_one(session, connector, job, provider, item, seen_hashes_this_job)
            if outcome == "INGESTED":
                job.files_ingested += 1
            elif outcome == "DUPLICATE":
                job.files_duplicate += 1
            elif outcome == "CHANGED":
                job.files_changed += 1
            else:
                job.files_failed += 1
            session.commit()  # per-file commit: progress polling sees live counts

        # Extraction turns parsed documents into ordinary CANDIDATE assets -
        # the same call the manual upload flow makes. A failure here is
        # recorded, never silent, but does not undo the ingested documents.
        auto_approval = None
        if auto_extract and job.files_ingested > 0:
            try:
                extraction.extract_knowledge_assets_from_project(session, job.project_id)
            except Exception as e:
                job.error = f"Asset extraction failed after ingestion: {e}"
                session.commit()
            else:
                # Policy-Based Auto Approval (MVP 0.10.2): only the documents
                # this job newly ingested. CHANGED files route through the
                # revision machinery and always wait for a human.
                try:
                    new_doc_ids = [r.document_id for r in session.query(db.SourceDocument).filter(
                        db.SourceDocument.ingestion_job_id == job.id,
                        db.SourceDocument.status == "INGESTED",
                        db.SourceDocument.document_id.isnot(None)).all()]
                    auto_approval = policy.apply_auto_approval(
                        session, job.project_id, new_doc_ids,
                        connector_id=connector.id, ingestion_job_id=job.id)
                except Exception as e:
                    job.error = f"Policy auto-approval failed after extraction: {e}"
                    session.commit()

        job.status = "COMPLETED"
        job.completed_at = datetime.datetime.utcnow()
        session.commit()
        crud.log_audit_event(
            session, actor=f"connector:{connector.name}", event_type="INGESTION_JOB_COMPLETED",
            target_id=str(job.id),
            details=json.dumps({
                "files_discovered": job.files_discovered,
                "files_ingested": job.files_ingested,
                "files_duplicate": job.files_duplicate,
                "files_changed": job.files_changed,
                "files_failed": job.files_failed,
                "extraction_error": job.error,
                "auto_approval": auto_approval,
            }))
    except Exception as e:
        job.status = "FAILED"
        job.error = str(e)
        job.completed_at = datetime.datetime.utcnow()
        session.commit()
        crud.log_audit_event(
            session, actor=f"connector:{connector.name}", event_type="INGESTION_JOB_FAILED",
            target_id=str(job.id),
            details=json.dumps({"error": str(e),
                                "files_discovered": job.files_discovered,
                                "files_ingested": job.files_ingested}))
    session.refresh(job)
    return job


def _ingest_one(session: Session, connector: db.SourceConnector, job: db.IngestionJob,
                provider: LocalFolderProvider, item: ConnectorItem, seen_hashes: set) -> str:
    """Process one discovered item; always writes a SourceDocument row.
    The provider fetches (at ingest time, not discovery time); the
    framework hashes - the sha256 of the fetched bytes is the only change
    verdict, provider metadata is recorded as described context."""
    record = db.SourceDocument(
        project_id=job.project_id,
        connector_id=connector.id,
        ingestion_job_id=job.id,
        source_uri=item.uri,
    )
    try:
        fetched = provider.fetch(item)
        data = fetched.data
        record.size_bytes = fetched.metadata.get("size_bytes")
        record.source_modified_at = fetched.metadata.get("modified_at")
        record.file_hash = hashlib.sha256(data).hexdigest()

        # Change detection (MVP 0.10.1): within a connector, the source URI is
        # the file's identity. Same URI + different content = CHANGED - routed
        # through the EXISTING revision machinery, never a duplicate, never a
        # second document. The per-scan SourceDocument rows are the version
        # history: every hash this URI has ever presented is retained.
        prior = session.query(db.SourceDocument).filter(
            db.SourceDocument.connector_id == connector.id,
            db.SourceDocument.source_uri == record.source_uri,
            db.SourceDocument.document_id.isnot(None),
        ).order_by(db.SourceDocument.id.desc()).first()
        if prior and prior.file_hash == record.file_hash:
            # Unchanged at this URI since the last scan - stays a duplicate
            # even if its document's content has since moved on via another
            # path's change (the inventory, not the document hash, is the
            # source of truth for tracked files).
            record.status = "DUPLICATE"
            record.error = f"Unchanged since last scan (document {prior.document_id})"
            record.document_id = prior.document_id
            session.add(record)
            return "DUPLICATE"
        if prior and prior.file_hash and prior.file_hash != record.file_hash:
            doc = session.query(db.Document).filter(db.Document.id == prior.document_id).first()
            if doc:
                seen_hashes.add(record.file_hash)
                return _apply_source_change(session, connector, job, record, doc, item, data)

        duplicate_of = None
        if record.file_hash in seen_hashes:
            duplicate_of = "another file in this scan"
        else:
            existing = session.query(db.Document).filter(
                db.Document.project_id == job.project_id,
                db.Document.content_hash == record.file_hash
            ).first()
            if existing:
                duplicate_of = f"document {existing.id} ({existing.filename})"
                record.document_id = existing.id

        if duplicate_of:
            record.status = "DUPLICATE"
            record.error = f"Identical content already ingested: {duplicate_of}"
            session.add(record)
            return "DUPLICATE"
        seen_hashes.add(record.file_hash)

        # From here on, the file becomes an ORDINARY document: same upload
        # dir, same create_document, same parser as the manual upload flow.
        # item.name is provider-supplied input - strip any path components
        # so no provider can stage outside UPLOAD_DIR (defense in depth).
        filename = os.path.basename((item.name or "").replace("\\", "/")).strip() or "unnamed"
        dest_path = os.path.join(UPLOAD_DIR, f"{job.project_id}_c{connector.id}_{record.file_hash[:8]}_{filename}")
        with open(dest_path, "wb") as out:
            out.write(data)

        doc = crud.create_document(
            session,
            project_id=job.project_id,
            filename=filename,
            file_path=dest_path,
            department="General",
            owner=f"connector:{connector.name}",
        )
        doc.content_hash = record.file_hash
        session.commit()

        ingestion.parse_and_index_document(session, doc.id)
        session.refresh(doc)
        if doc.status == "FAILED":
            # The SourceDocument row records the failure; remove the dead
            # Document so the inventory stays clean and a rescan can retry
            # the file instead of deduping against a failed ingest.
            session.delete(doc)
            session.commit()
            try:
                os.remove(dest_path)
            except OSError:
                pass
            record.status = "FAILED"
            record.error = "Document parsing failed (unsupported or corrupt file content)"
            session.add(record)
            return "FAILED"
        record.document_id = doc.id

        record.status = "INGESTED"
        session.add(record)
        return "INGESTED"
    except Exception as e:
        record.status = "FAILED"
        record.error = str(e)
        session.add(record)
        return "FAILED"


def _extract_text(file_path: str) -> str:
    """Text extraction for changed files - mirrors the ingestion fallbacks
    (.txt/.md read, .pdf via pypdf, .docx via python-docx)."""
    lower = file_path.lower()
    if lower.endswith(".docx"):
        import docx
        d = docx.Document(file_path)
        parts = [p.text for p in d.paragraphs if p.text.strip()]
        for table in d.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n\n".join(parts)
    if lower.endswith(".pdf"):
        import pypdf
        reader = pypdf.PdfReader(file_path)
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def _discard_temp_asset(session: Session, asset: db.KnowledgeAsset):
    """Temp extraction rows that reconciled into existing assets are removed
    with their quality scores - they were never reviewable objects."""
    session.query(db.QualityScore).filter(db.QualityScore.asset_id == asset.id).delete(synchronize_session=False)
    session.delete(asset)


def _apply_source_change(session: Session, connector: db.SourceConnector, job: db.IngestionJob,
                         record: db.SourceDocument, doc: db.Document,
                         item: ConnectorItem, data: bytes) -> str:
    """A source file changed (MVP 0.10.1): re-extract the new content and
    reconcile it against the document's existing assets through the EXISTING
    governance machinery -
      approved asset, content differs   -> candidate revision (Revision Workbench)
      non-approved asset, differs       -> edited in place (the platform rule)
      no matching asset                 -> new CANDIDATE asset
      asset with no counterpart anymore -> reported as possibly_stale, untouched
    Old chunks are kept so approved assets retain verifiable provenance to the
    content they were approved against until their revisions are approved."""
    try:
        # item.name, not URI parsing - deriving a display name from a URI
        # is provider knowledge (identical bytes for local paths). Sanitized
        # like the ingest path: provider-supplied names never carry paths.
        filename = os.path.basename((item.name or "").replace("\\", "/")).strip() or "unnamed"
        dest_path = os.path.join(UPLOAD_DIR, f"{job.project_id}_c{connector.id}_{record.file_hash[:8]}_{filename}")
        with open(dest_path, "wb") as out:
            out.write(data)
        new_text = _extract_text(dest_path)
        if not new_text.strip():
            raise ValueError("Changed file produced no extractable text")

        # Same Document row, updated source facts - never a second document.
        doc.file_path = dest_path
        doc.content_hash = record.file_hash
        doc.modified_at = datetime.datetime.utcnow()
        session.commit()

        next_index = session.query(db.DocumentChunk).filter(
            db.DocumentChunk.document_id == doc.id).count()
        chunk = db.DocumentChunk(document_id=doc.id, text=new_text, chunk_index=next_index)
        session.add(chunk)
        session.commit()
        session.refresh(chunk)

        # Scoped re-extraction of just the new content, with the same
        # extractors first ingestion used.
        before_ids = {a.id for a in session.query(db.KnowledgeAsset).filter(
            db.KnowledgeAsset.document_id == doc.id).all()}
        api_key = os.environ.get("OPENAI_API_KEY")
        if api_key and not api_key.startswith("mock-"):
            extraction.extract_via_llm(session, job.project_id, doc, chunk, api_key)
        else:
            extraction.extract_via_rules(session, job.project_id, doc, chunk)

        all_assets = session.query(db.KnowledgeAsset).filter(
            db.KnowledgeAsset.document_id == doc.id).all()
        fresh = [a for a in all_assets if a.id not in before_ids]
        existing = [a for a in all_assets if a.id in before_ids]

        summary = {"revisions_created": [], "updated_in_place": 0, "assets_added": 0,
                   "unchanged": 0, "skipped_pending_review": [], "possibly_stale": []}
        matched_ids = set()
        for new_asset in fresh:
            match = next((a for a in existing
                          if a.type == new_asset.type and a.name == new_asset.name
                          and a.id not in matched_ids), None)
            if not match:
                summary["assets_added"] += 1
                continue  # genuinely new knowledge stays as an ordinary CANDIDATE
            matched_ids.add(match.id)
            if (match.content or "").strip() == (new_asset.content or "").strip():
                summary["unchanged"] += 1
                _discard_temp_asset(session, new_asset)
                continue
            if match.status == "APPROVED":
                try:
                    rev = revisions.create_candidate_revision(
                        session, match.id, new_asset.content,
                        actor=f"connector:{connector.name}",
                        change_reason=f"Source file changed: {record.source_uri}")
                    summary["revisions_created"].append({"asset_id": match.id, "revision_id": rev.id})
                except ValueError as e:
                    # One pending candidate per asset (strictly linear) -
                    # reported, never silent. The next scan retries.
                    summary["skipped_pending_review"].append({"asset_id": match.id, "reason": str(e)})
            else:
                match.content = new_asset.content
                match.chunk_id = new_asset.chunk_id
                match.source_hash = new_asset.source_hash
                summary["updated_in_place"] += 1
            _discard_temp_asset(session, new_asset)
        summary["possibly_stale"] = [
            {"asset_id": a.id, "name": a.name} for a in existing if a.id not in matched_ids]
        session.commit()

        record.status = "CHANGED"
        record.document_id = doc.id
        record.details_json = json.dumps(summary)
        session.add(record)
        crud.log_audit_event(
            session, actor=f"connector:{connector.name}", event_type="SOURCE_CHANGE_DETECTED",
            target_id=str(doc.id),
            details=json.dumps({"source_uri": record.source_uri, "new_hash": record.file_hash, **summary}))
        return "CHANGED"
    except Exception as e:
        record.status = "FAILED"
        record.error = f"Change handling failed: {e}"
        session.add(record)
        return "FAILED"
