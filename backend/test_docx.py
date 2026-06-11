import os
import sys

# Ensure backend folder is in path
sys.path.append(os.path.join(os.path.dirname(__file__)))

import docx
from app import database as db
from app import schemas
from app import crud
from app import ingestion
import test_support

def test_docx_ingestion():
    print("Initializing test database for DOCX review...")
    db.init_db()
    session = db.SessionLocal()
    
    try:
        # Create a real docx document
        docx_path = "./uploads/real_clinical_sop.docx"
        os.makedirs("./uploads", exist_ok=True)
        
        print("Compiling real Word document using python-docx...")
        doc_writer = docx.Document()
        doc_writer.add_heading('SOP-777: Biological Specimen Handling', 0)
        
        doc_writer.add_paragraph(
            "Purpose: Define the governance protocol for biological specimen logistics.\n"
            "Department: Clinical Lab Operations"
        )
        
        doc_writer.add_paragraph(
            "Policy: All specimens must be logged into the temperature audit inventory within 15 minutes of intake. "
            "Any deviation from temperature parameters requires immediate escalation to the Operations Director."
        )
        
        # Add a table
        table = doc_writer.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Role'
        hdr_cells[1].text = 'Responsibility'
        
        row_cells = table.add_row().cells
        row_cells[0].text = 'Lab Technician'
        row_cells[1].text = 'Log specimens on intake'
        
        doc_writer.save(docx_path)
        print(f"Word document saved to {docx_path}")
        
        # Log Document
        customer = crud.get_or_create_default_customer(session)
        actor = test_support.governed_actor(session)
        project = crud.create_project(session, schemas.ProjectCreate(
            name="DOCX Verification Project",
            description="Testing DOCX parser",
            customer_id=customer.id
        ), actor=actor)

        db_doc = crud.create_document(
            session,
            project_id=project.id,
            filename="real_clinical_sop.docx",
            file_path=docx_path,
            department="Clinical Labs",
            owner="Lab Director",
            actor=actor
        )
        
        print("\nTriggering parse pipeline...")
        success = ingestion.parse_and_index_document(session, db_doc.id)
        print(f"Parse Success: {success}")
        
        session.refresh(db_doc)
        print(f"Document parsing status: {db_doc.status}")
        
        if success:
            chunks = crud.get_document_chunks(session, db_doc.id)
            print(f"\nExtracted Chunks ({len(chunks)}):")
            for c in chunks:
                print(f"--- Chunk {c.chunk_index} ---")
                print(c.text)
                print("---------------------")
                # Assertions
                assert "SOP-777" in c.text or "specimen" in c.text or "Role" in c.text
                assert "[Content_Types].xml" not in c.text
                assert "word/document.xml" not in c.text
            print("\nVerification PASSED: Extracted human-readable text correctly. No raw OOXML found.")
        else:
            print("\nVerification FAILED: Parser failed to ingest.")
            sys.exit(1)
            
    finally:
        session.close()

if __name__ == "__main__":
    test_docx_ingestion()
